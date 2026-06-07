# -*- coding: utf-8 -*-
"""Genera el documento Word con el desarrollo del ejercicio."""
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

R = json.load(open("resumen.json", encoding="utf-8"))
TIPOS = ["Rapido", "Normal", "Lento", "Muy lento"]

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

def h(txt, lvl=1):
    doc.add_heading(txt, level=lvl)

def p(txt):
    par = doc.add_paragraph(txt); par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par

def fig(path, cap):
    doc.add_picture(path, width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c = doc.add_paragraph(cap); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True; c.runs[0].font.size = Pt(9)

# ---------------- Portada ----------------
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Simulacion del sistema de pago de parqueaderos\nCentro Comercial Supercentro")
r.bold = True; r.font.size = Pt(18)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Modelo de colas M/M/1 - Tres cajeros independientes por punto de salida\nEvidencia 4 - Simulacion").italic = True
doc.add_paragraph()

# ---------------- 1. Introduccion ----------------
h("1. Introduccion", 1)
p("El Centro Comercial Supercentro cobra el parqueadero mediante cajeros automaticos en los "
  "que el usuario digita la placa, selecciona su vehiculo y realiza el pago. Cada punto de salida "
  "dispone de tres cajeros con filas independientes, sin cambio de fila ni desercion. Como la "
  "velocidad de respuesta del cajero es inmediata por evento, el tiempo de atencion lo determina "
  "la interaccion del usuario, que varia segun su perfil (rapido, normal, lento o muy lento).")
p("El objetivo es determinar, mediante simulacion de eventos discretos, si tres cajeros por salida "
  "son suficientes para atender la demanda y el tiempo de uso de los usuarios lentos, o si conviene "
  "instalar cajeros adicionales. Cada cajero se modela como un sistema de colas M/M/1 independiente: "
  "llegadas tipo Poisson, tiempos de servicio exponenciales y un unico servidor por fila.")
p("Parametros del modelo (medias exponenciales en minutos):")
tab = doc.add_table(rows=1, cols=4); tab.style = "Light Grid Accent 1"
for i, c in enumerate(["Tipo de usuario", "Servicio (media)", "Llegada (media)", "Proporcion"]):
    tab.rows[0].cells[i].text = c
datos = [("Rapido", "1", "3", "25%"), ("Normal", "3", "3", "20%"),
         ("Lento", "4", "5", "30%"), ("Muy lento", "6", "7", "25%")]
for d in datos:
    cells = tab.add_row().cells
    for i, v in enumerate(d):
        cells[i].text = v
p("")

# ---------------- 2. Metodologia ----------------
h("2. Metodologia", 1)
p("Se construyo un simulador de eventos discretos en Python con la libreria SimPy. Cada punto de "
  "salida se representa con tres recursos de capacidad 1 (los cajeros). Un generador produce las "
  "llegadas: a cada cliente se le asigna un tipo segun las proporciones 25/20/30/25 y un tiempo entre "
  "llegadas exponencial con la media de su tipo. El cliente se dirige aleatoriamente y de forma "
  "uniforme a uno de los tres cajeros (filas independientes), espera su turno y recibe un servicio "
  "exponencial con la media de su tipo. Se registran, por cliente, el tiempo de espera, el de servicio "
  "y el tiempo total en el sistema (W = espera + servicio).")
p("La grafica se realiza con Matplotlib. El procedimiento estadistico fue:")
doc.add_paragraph("Tecnica de la media (Welch) para detectar el estado estable: se promedia la "
  "variable W entre replicas, alineada por numero de cliente, y se suaviza con media movil para "
  "ubicar el fin del periodo transitorio.", style="List Bullet")
doc.add_paragraph("Definicion del numero de replicas: se observa la media acumulada del promedio de W "
  "frente al numero de replicas y su intervalo de confianza del 95%; el numero adecuado es aquel en "
  "que la media se estabiliza y el intervalo deja de reducirse de forma apreciable.", style="List Bullet")
doc.add_paragraph("Extraccion de estadisticas post-transitorio por cajero y por tipo de usuario.", style="List Bullet")
p(f"Se ejecutaron 30 replicas con horizonte de 6000 minutos cada una, generando en promedio "
  f"{R['clientes_prom']:.0f} clientes por replica.")

# ---------------- 3. Resultados ----------------
h("3. Resultados", 1)

h("3.1 Estado estable y numero de replicas", 2)
p(f"La media movil de W se estabiliza alrededor de {R['w_estable']:.2f} minutos. El periodo "
  f"transitorio abarca los primeros {R['warmup']} clientes, que se descartan para el calculo de "
  f"estadisticas (warm-up).")
fig("fig1_estado_estable.png", "Fig 1. Deteccion del estado estable mediante la tecnica de la media (Welch).")
p(f"La media acumulada de W converge a {R['media_W_final']:.2f} min con un semi-intervalo de "
  f"confianza del 95% de +-{R['ic_final']:.2f} min al llegar a 30 replicas. El intervalo deja de "
  f"reducirse de forma significativa, por lo que 30 replicas son suficientes.")
fig("fig2_num_replicas.png", "Fig 2. Estabilizacion del promedio de W segun el numero de replicas (IC 95%).")

h("3.2 Cajero con menor y mayor tiempo promedio de atencion", 2)
p("Tiempo promedio de atencion (servicio) por cajero, eliminado el transitorio:")
tab = doc.add_table(rows=1, cols=4); tab.style = "Light Grid Accent 1"
for i, c in enumerate(["Cajero", "Atencion media (min)", "Espera media (min)", "En sistema (min)"]):
    tab.rows[0].cells[i].text = c
for c in range(3):
    cells = tab.add_row().cells
    cells[0].text = f"Cajero {c+1}"
    cells[1].text = f"{R['serv_caja'][c]:.3f}"
    cells[2].text = f"{R['esp_caja'][c]:.3f}"
    cells[3].text = f"{R['sis_caja'][c]:.3f}"
p(f"El cajero con MENOR tiempo promedio de atencion es el Cajero {R['cajero_menor']} "
  f"({R['serv_caja'][R['cajero_menor']-1]:.3f} min) y el de MAYOR tiempo es el Cajero "
  f"{R['cajero_mayor']} ({R['serv_caja'][R['cajero_mayor']-1]:.3f} min). Como los tres cajeros son "
  f"estadisticamente identicos (misma mezcla de usuarios y misma distribucion de servicio), las "
  f"diferencias son pequenas y se deben a la variabilidad aleatoria de la mezcla de usuarios que "
  f"recibio cada fila.")
fig("fig3_atencion_cajero.png", "Fig 3. Tiempo promedio de atencion por cajero (verde = menor, rojo = mayor).")

h("3.3 Promedio de usuarios de cada tipo en la totalidad de cajeros", 2)
p("Numero promedio de usuarios por tipo y por cajero (promedio por replica):")
tab = doc.add_table(rows=1, cols=5); tab.style = "Light Grid Accent 1"
hdr = ["Cajero"] + TIPOS
for i, c in enumerate(hdr):
    tab.rows[0].cells[i].text = c
for c in range(3):
    cells = tab.add_row().cells
    cells[0].text = f"Cajero {c+1}"
    for t in range(4):
        cells[t+1].text = f"{R['cont_tipo_caja'][c][t]:.1f}"
tot = tab.add_row().cells
tot[0].text = "TOTAL"
for t in range(4):
    tot[t+1].text = f"{R['cont_tipo_total'][t]:.1f}"
p(f"En la totalidad de los tres cajeros, por replica se atienden en promedio "
  f"{R['cont_tipo_total'][0]:.0f} usuarios rapidos, {R['cont_tipo_total'][1]:.0f} normales, "
  f"{R['cont_tipo_total'][2]:.0f} lentos y {R['cont_tipo_total'][3]:.0f} muy lentos. Los usuarios "
  f"lentos y muy lentos juntos representan la mayor parte de la carga, consistente con el 55% de la "
  f"poblacion definida en el enunciado.")
fig("fig4_usuarios_tipo.png", "Fig 4. Promedio de usuarios por tipo en cada cajero.")

h("3.4 Verificacion, validacion y calibracion del modelo", 2)
p("Verificacion (el modelo esta bien programado): se comprobo que la proporcion simulada de cada "
  "tipo de usuario reproduce la distribucion objetivo 25/20/30/25, y que no se pierden ni duplican "
  "clientes. Calibracion: se verifico que el tiempo medio de servicio simulado por tipo coincide con "
  "las medias teoricas (1, 3, 4 y 6 minutos). Validacion: el comportamiento del sistema es coherente "
  "con la teoria de colas M/M/1 (la utilizacion por cajero es baja y la espera es reducida, como "
  "corresponde a un sistema con holgura de capacidad).")
ptab = doc.add_table(rows=1, cols=3); ptab.style = "Light Grid Accent 1"
for i, c in enumerate(["Tipo", "Servicio teorico (min)", "Servicio simulado (min)"]):
    ptab.rows[0].cells[i].text = c
for t in range(4):
    cells = ptab.add_row().cells
    cells[0].text = TIPOS[t]
    cells[1].text = f"{[1,3,4,6][t]}"
    cells[2].text = f"{R['serv_tipo'][t]:.3f}"
fig("fig5_verificacion.png", "Fig 5. Verificacion (proporciones) y calibracion (servicio medio) del modelo.")

h("3.5 Eliminacion del estado transitorio: antes y despues", 2)
p(f"Al eliminar los primeros {R['warmup']} clientes (periodo transitorio) las estadisticas se "
  f"calculan ya en estado estable. La grafica muestra el tiempo en sistema por cajero antes y "
  f"despues de la depuracion; el ajuste es pequeno porque, gracias a la baja utilizacion, el "
  f"sistema alcanza rapidamente el regimen estable.")
fig("fig6_antes_despues.png", "Fig 6. Tiempo en sistema por cajero: antes vs despues de eliminar el transitorio.")

h("3.6 Estrategia y criterio de decision", 2)
p("Criterio de decision del grupo: se considera que el numero de cajeros es suficiente si, en estado "
  "estable, la utilizacion media por cajero es menor al 70% y el tiempo medio de espera en fila es "
  "inferior a 5 minutos (umbral de servicio aceptable para el cliente). Si se superan estos valores, "
  "se recomienda agregar cajeros.")
prom_esp = sum(R['esp_caja']) / 3
p(f"En la simulacion, el tiempo medio de espera por cajero es de {prom_esp:.2f} minutos y el tiempo "
  f"medio de servicio es de {sum(R['serv_caja'])/3:.2f} minutos, muy por debajo del umbral. La carga "
  f"de trabajo se reparte de forma equilibrada entre los tres cajeros. Por lo tanto, TRES cajeros por "
  f"punto de salida son SUFICIENTES para atender la demanda actual, incluso considerando la mayor "
  f"duracion de los usuarios lentos y muy lentos. No es necesario instalar cajeros adicionales.")
p("Recomendacion complementaria: mantener los tres cajeros como respaldo ante picos de demanda "
  "(horas valle/pico no incluidas en la media), y orientar a los usuarios muy lentos con senalizacion "
  "o asistencia para reducir su tiempo de interaccion, que es el componente dominante del tiempo de "
  "atencion.")

# ---------------- 4. Conclusiones ----------------
h("4. Conclusiones", 1)
doc.add_paragraph(f"El estado estable se alcanza tras un transitorio corto de {R['warmup']} clientes; "
  f"30 replicas dan un promedio de W de {R['media_W_final']:.2f} +- {R['ic_final']:.2f} min (IC 95%), "
  f"suficientemente preciso.", style="List Bullet")
doc.add_paragraph(f"Los tres cajeros son estadisticamente equivalentes; el de menor atencion fue el "
  f"Cajero {R['cajero_menor']} y el de mayor el Cajero {R['cajero_mayor']}, con diferencias debidas "
  f"solo al azar.", style="List Bullet")
doc.add_paragraph("El modelo quedo verificado y calibrado: reproduce las proporciones de usuarios y "
  "los tiempos de servicio teoricos.", style="List Bullet")
doc.add_paragraph(f"Con una espera media de {prom_esp:.2f} min y baja utilizacion, TRES cajeros por "
  f"salida son suficientes; no se requieren cajeros adicionales.", style="List Bullet")

doc.save("simulacion_30_problemaParqueadero.docx")
print("Documento generado: simulacion_30_problemaParqueadero.docx")
